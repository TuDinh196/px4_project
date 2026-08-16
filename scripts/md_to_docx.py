#!/usr/bin/env python3
"""
Convert Markdown to DOCX
========================
Converts all Markdown (.md) documentation files in the project into formatted
Microsoft Word (.docx) documents, preserving headings, tables, code blocks,
bullet lists, and embedded images.
"""

import re
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    """Set shading/background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def convert_md_to_docx(md_path: Path, output_docx_path: Path):
    """Parse a Markdown file and convert it into a formatted Word (.docx) document."""
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base directory for relative image paths
    base_dir = md_path.parent

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal in_table, table_lines
        if not table_lines:
            in_table = False
            return

        # Parse table rows
        rows_data = []
        for line in table_lines:
            line_str = line.strip()
            # Skip separator rows like |:---|:---| or |---|---|
            if not line_str or "---" in line_str:
                continue
            cells = [c.strip() for c in line_str.split("|")[1:-1]]
            if cells:
                rows_data.append(cells)

        if rows_data:
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_text
                        # Style header row
                        if r_idx == 0:
                            set_cell_background(cell, "1E293B")
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(248, 250, 252)
                        else:
                            bg_color = "F1F5F9" if r_idx % 2 == 1 else "FFFFFF"
                            set_cell_background(cell, bg_color)
            doc.add_paragraph()  # Spacing

        table_lines = []
        in_table = False

    def process_inline_formatting(paragraph, text):
        """Add runs to paragraph handling **bold**, *italic*, `code`, and images."""
        # Check for image ![caption](path)
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)", text.strip())
        if img_match:
            caption, img_rel_path = img_match.group(1), img_match.group(2)
            img_abs_path = (base_dir / img_rel_path).resolve()
            if img_abs_path.exists():
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                try:
                    run_img.add_picture(str(img_abs_path), width=Inches(5.8))
                except Exception:
                    run_img.add_picture(str(img_abs_path))
                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run(f"Hinh: {caption}")
                    run_cap.font.italic = True
                    run_cap.font.size = Pt(9.5)
                    run_cap.font.color.rgb = RGBColor(100, 116, 139)
                return True
            else:
                paragraph.add_run(f"[Hinh anh: {caption} ({img_rel_path})]")
                return True

        # Parse bold/inline code formatting
        parts = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(225, 29, 72)
            else:
                paragraph.add_run(part)
        return False

    i = 0
    while i < len(lines):
        line = lines[i]
        raw_line = line.rstrip("\n")

        # Code block toggle ```
        if raw_line.strip().startswith("```"):
            if in_code_block:
                # End of code block
                in_code_block = False
                code_text = "\n".join(code_block_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "0F172A")
                p = cell.paragraphs[0]
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(241, 245, 249)
                doc.add_paragraph()  # spacing
                code_block_lines = []
            else:
                if in_table:
                    flush_table()
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(raw_line)
            i += 1
            continue

        # Table rows
        if raw_line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(raw_line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Headings
        if raw_line.startswith("# "):
            p = doc.add_heading(level=1)
            process_inline_formatting(p, raw_line[2:].strip())
            p.style.font.color.rgb = RGBColor(15, 23, 42)
        elif raw_line.startswith("## "):
            p = doc.add_heading(level=2)
            process_inline_formatting(p, raw_line[3:].strip())
            p.style.font.color.rgb = RGBColor(30, 41, 59)
        elif raw_line.startswith("### "):
            p = doc.add_heading(level=3)
            process_inline_formatting(p, raw_line[4:].strip())
            p.style.font.color.rgb = RGBColor(51, 65, 85)
        elif raw_line.startswith("#### "):
            p = doc.add_heading(level=4)
            process_inline_formatting(p, raw_line[5:].strip())
        elif raw_line.startswith("- ") or raw_line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            process_inline_formatting(p, raw_line[2:].strip())
        elif re.match(r"^\d+\.\s+", raw_line):
            text = re.sub(r"^\d+\.\s+", "", raw_line)
            p = doc.add_paragraph(style="List Number")
            process_inline_formatting(p, text.strip())
        elif raw_line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run("  ")
            process_inline_formatting(p, raw_line[2:].strip())
            for r in p.runs:
                r.font.italic = True
                r.font.color.rgb = RGBColor(71, 85, 105)
        elif raw_line.strip() == "---":
            # Horizontal rule
            p = doc.add_paragraph()
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(203, 213, 225)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif raw_line.strip():
            p = doc.add_paragraph()
            process_inline_formatting(p, raw_line.strip())

        i += 1

    if in_table:
        flush_table()

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))
    print(f"  [OK] Converted: {md_path.name} -> {output_docx_path.name}")


def main():
    project_root = Path(__file__).resolve().parents[1]
    docs_dir = project_root / "docs"

    print("=" * 60)
    print("  Convert All Markdown (.md) to Word (.docx) Documents")
    print("=" * 60)

    md_files = list(docs_dir.glob("*.md")) + list(project_root.glob("*.md"))

    if not md_files:
        print("[FAIL] No Markdown (.md) files found")
        return

    converted_count = 0
    for md_file in sorted(md_files):
        docx_file = md_file.with_suffix(".docx")
        try:
            convert_md_to_docx(md_file, docx_file)
            converted_count += 1
        except Exception as e:
            print(f"[FAIL] Failed to convert {md_file.name}: {e}")

    print("\n" + "=" * 60)
    print(f"  [DONE] Conversion Complete: {converted_count}/{len(md_files)} files converted!")
    print("=" * 60)


if __name__ == "__main__":
    main()
