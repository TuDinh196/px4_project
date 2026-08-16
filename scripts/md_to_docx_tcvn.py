#!/usr/bin/env python3
"""
Convert Markdown to DOCX — Vietnamese Standard (TCVN / Nghị định 30/2020/NĐ-CP)
================================================================================
Formats Word (.docx) documents according to official Vietnamese document standards:
  - Font: Times New Roman (Unicode)
  - Margins: Left 3.0 cm, Right 2.0 cm, Top 2.5 cm, Bottom 2.5 cm
  - Title/Headings: 14-16pt Bold, Dark Blue palette (#002060, #1F4E78, #2F5597)
  - Body Text: 13pt Regular, Justified alignment, 1.25 line spacing, 1.27cm first line indent
  - Math Equations: Formatted in Cambria Math / Italic with clean spacing & Unicode notation
  - Tables: Styled header (#1F4E78), alternating row shading (#F2F5F9), centered alignment
  - Images & Captions: Centered with italic 10pt Slate caption
"""

import re
from pathlib import Path

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    """Set shading/background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in twips (1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def clean_latex_math(text):
    """Transforms LaTeX math syntax into elegant, highly readable Unicode mathematical notation."""
    if not text:
        return ""

    def replace_frac(m):
        num, den = m.group(1), m.group(2)
        return f"({num}/{den})"

    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", replace_frac, text)
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", replace_frac, text)

    # Handle \dot{x}, \ddot{x}
    text = re.sub(r"\\ddot\{([a-zA-Z])\}", r"\1̈", text)
    text = re.sub(r"\\dot\{([a-zA-Z])\}", r"\1̇", text)
    text = re.sub(r"\\dot\{\\phi\}", "φ̇", text)

    # Greek letters
    greek = [
        (r"\phi", "φ"), (r"\theta", "θ"), (r"\psi", "ψ"),
        (r"\omega", "ω"), (r"\alpha", "α"), (r"\beta", "β"),
        (r"\delta", "δ"), (r"\rho", "ρ"), (r"\eta", "η"),
        (r"\tau", "τ"), (r"\sigma", "σ"), (r"\gamma", "γ"),
        (r"\lambda", "λ"), (r"\mu", "μ"), (r"\nu", "ν"),
        (r"\Delta", "Δ"), (r"\Omega", "Ω"), (r"\Phi", "Φ"),
        (r"\Theta", "Θ"), (r"\Psi", "Ψ"), (r"\Sigma", "Σ")
    ]
    for lat, sym in greek:
        text = re.sub(lat.replace("\\", r"\\") + r"(?![a-zA-Z])", sym, text)

    # Operations & Symbols
    symbols = [
        (r"\times", "×"), (r"\cdot", "·"), (r"\approx", "≈"),
        (r"\le", "≤"), (r"\ge", "≥"), (r"\leq", "≤"), (r"\geq", "≥"),
        (r"\ne", "≠"), (r"\neq", "≠"), (r"\pm", "±"), (r"\mp", "∓"),
        (r"\in", "∈"), (r"\to", "→"), (r"\rightarrow", "→"),
        (r"\leftarrow", "←"), (r"\leftrightarrow", "↔"),
        (r"\infty", "∞"), (r"\nabla", "∇"), (r"\partial", "∂"),
        (r"\sum", "∑"), (r"\int", "∫"), (r"\sqrt", "√"),
        (r"\circ", "°"), (r"\\", " "), (r"\,", " "), (r"\;", " "),
        (r"\quad", "  "), (r"\qquad", "    "),
        (r"\_", "_"), (r"\%", "%")
    ]
    for lat, sym in symbols:
        text = re.sub(lat.replace("\\", r"\\"), sym, text)

    # Subscripts and superscripts
    sub_map = str.maketrans("0123456789aehijklmnoprstuvx", "₀₁₂₃₄₅₆₇₈₉ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
    sup_map = str.maketrans("0123456789+-=()nTi", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿᵀⁱ")

    def replace_sub(m):
        return m.group(1).translate(sub_map)

    def replace_sup(m):
        return m.group(1).translate(sup_map)

    text = re.sub(r"_\{([0-9a-zA-Z_]+)\}", replace_sub, text)
    text = re.sub(r"\^\{([0-9a-zA-Z_+\-()]+)\}", replace_sup, text)
    text = re.sub(r"_([0-9a-zA-Z])", replace_sub, text)
    text = re.sub(r"\^([0-9a-zA-Z])", replace_sup, text)

    return text.strip()


def convert_md_to_docx_tcvn(md_path: Path, output_docx_path: Path):
    """Converts a Markdown file to an official TCVN-formatted Microsoft Word document."""
    base_dir = md_path.parent
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal in_table, table_lines
        if not table_lines:
            in_table = False
            return

        rows_data = []
        for line_str in table_lines:
            if re.match(r"^\s*\|?\s*[-:]+[-| :]*\|?\s*$", line_str):
                continue
            cells = [c.strip() for c in line_str.split("|")[1:-1]]
            if cells:
                rows_data.append(cells)

        if rows_data:
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                        cell_clean = clean_latex_math(cell_text)
                        cell.text = cell_clean
                        if r_idx == 0:
                            set_cell_background(cell, "1F4E78")
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in p.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(11)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            bg_color = "F2F5F9" if r_idx % 2 == 1 else "FFFFFF"
                            set_cell_background(cell, bg_color)
                            for p in cell.paragraphs:
                                is_num = cell_clean.replace('.', '').replace('-', '').isdigit()
                                if len(cell_clean) < 12 or is_num:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in p.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(15, 23, 42)
            doc.add_paragraph()
        table_lines = []
        in_table = False

    def process_inline_runs(paragraph, text, default_font="Times New Roman", default_size=13.0):
        """Process inline text formatting: **bold**, *italic*, `code`, and $math$."""
        # Image check ![caption](path)
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)", text.strip())
        if img_match:
            caption, img_rel_path = img_match.group(1), img_match.group(2)
            img_abs_path = (base_dir / img_rel_path).resolve()
            if img_abs_path.exists():
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(2)
                run_img = p_img.add_run()
                try:
                    run_img.add_picture(str(img_abs_path), width=Inches(5.6))
                except Exception:
                    run_img.add_picture(str(img_abs_path))

                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(8)
                    run_cap = p_cap.add_run(f"Hinh: {caption}")
                    run_cap.font.name = "Times New Roman"
                    run_cap.font.italic = True
                    run_cap.font.size = Pt(10)
                    run_cap.font.color.rgb = RGBColor(71, 85, 105)
                return True
            else:
                paragraph.add_run(f"[Hinh: {caption}]")
                return True

        # Split inline math $...$, bold **...**, code `...`, italic *...*
        parts = re.split(r"(\$\$.*?\$\$|\$.*?\$|\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("$$") and part.endswith("$$"):
                # Display math inside line
                clean = clean_latex_math(part[2:-2])
                run = paragraph.add_run(f" {clean} ")
                run.font.name = "Cambria Math"
                run.font.size = Pt(default_size)
                run.font.italic = True
                run.font.color.rgb = RGBColor(30, 41, 59)
            elif part.startswith("$") and part.endswith("$"):
                # Inline math
                clean = clean_latex_math(part[1:-1])
                run = paragraph.add_run(clean)
                run.font.name = "Cambria Math"
                run.font.size = Pt(default_size)
                run.font.italic = True
                run.font.color.rgb = RGBColor(30, 41, 59)
            elif part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.font.name = default_font
                run.font.size = Pt(default_size)
                run.font.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42)
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = default_font
                run.font.size = Pt(default_size)
                run.font.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(194, 24, 91)
            else:
                run = paragraph.add_run(part)
                run.font.name = default_font
                run.font.size = Pt(default_size)
                run.font.color.rgb = RGBColor(15, 23, 42)
        return False

    i = 0
    while i < len(lines):
        line = lines[i]
        raw_line = line.rstrip("\n")

        # Code Block Toggle ```
        if raw_line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                code_text = "\n".join(code_block_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                # Light background (#F1F5F9) for formal document printing
                set_cell_background(cell, "F1F5F9")
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(15, 23, 42)  # Dark slate text for high contrast
                doc.add_paragraph()  # Spacing
                code_block_lines = []
            else:
                if in_table:
                    flush_table()
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(raw_line)
            i += 1
            continue

        # Standalone Display Math Block $$ ... $$
        if (
            raw_line.strip().startswith("$$")
            and raw_line.strip().endswith("$$")
            and len(raw_line.strip()) > 4
        ):
            if in_table:
                flush_table()
            p_math = doc.add_paragraph()
            p_math.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_math.paragraph_format.space_before = Pt(6)
            p_math.paragraph_format.space_after = Pt(6)
            math_clean = clean_latex_math(raw_line.strip()[2:-2])
            run_math = p_math.add_run(math_clean)
            run_math.font.name = "Cambria Math"
            run_math.font.size = Pt(13.0)
            run_math.font.italic = True
            run_math.font.bold = True
            run_math.font.color.rgb = RGBColor(30, 41, 59)
            i += 1
            continue

        # Table rows
        if raw_line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(raw_line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Headings (TCVN Standard Colors & Sizes)
        if raw_line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = True
            process_inline_runs(
                p, raw_line[2:].strip().upper(),
                default_font="Times New Roman", default_size=16.0
            )
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0, 32, 96)  # Deep Navy Title #002060
        elif raw_line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            process_inline_runs(
                p, raw_line[3:].strip(),
                default_font="Times New Roman", default_size=14.0
            )
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(31, 78, 120)  # Section Header #1F4E78
        elif raw_line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            process_inline_runs(
                p, raw_line[4:].strip(),
                default_font="Times New Roman", default_size=13.0
            )
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(47, 85, 151)  # Subsection Header #2F5597
        elif raw_line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            process_inline_runs(
                p, raw_line[5:].strip(),
                default_font="Times New Roman", default_size=13.0
            )
            for r in p.runs:
                r.font.bold = True
        # Bullet Lists
        elif raw_line.startswith("- ") or raw_line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            process_inline_runs(
                p, raw_line[2:].strip(),
                default_font="Times New Roman", default_size=13.0
            )
        # Numbered Lists
        elif re.match(r"^\d+\.\s+", raw_line):
            text = re.sub(r"^\d+\.\s+", "", raw_line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            process_inline_runs(
                p, text.strip(),
                default_font="Times New Roman", default_size=13.0
            )
        # Blockquotes / Notes
        elif raw_line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.add_run("📌 ")
            process_inline_runs(
                p, raw_line[2:].strip(),
                default_font="Times New Roman", default_size=12.0
            )
            for r in p.runs:
                r.font.italic = True
                r.font.color.rgb = RGBColor(71, 85, 105)
        # Horizontal Rule
        elif raw_line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 45)
            run.font.color.rgb = RGBColor(203, 213, 225)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Regular Body Paragraphs (TCVN Standard)
        elif raw_line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Căn đều 2 bên
            p.paragraph_format.line_spacing = 1.25      # Giãn dòng 1.25
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Cm(1.27)  # Thụt lùi đầu dòng 1.27 cm
            process_inline_runs(
                p, raw_line.strip(),
                default_font="Times New Roman", default_size=13.0
            )

        i += 1

    if in_table:
        flush_table()

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))
    print(f"  [TCVN OK] {md_path.name} -> {output_docx_path.name}")


def main():
    project_root = Path(__file__).resolve().parents[1]
    docs_dir = project_root / "docs"

    print("=" * 65)
    print("  Convert Markdown to TCVN Word (.docx) - Nghi dinh 30/2020/ND-CP")
    print("=" * 65)

    md_files = sorted(list(docs_dir.glob("*.md")) + list(project_root.glob("*.md")))

    if not md_files:
        print("[FAIL] No Markdown (.md) files found")
        return

    converted_count = 0
    for md_file in md_files:
        docx_file = md_file.with_suffix(".docx")
        try:
            convert_md_to_docx_tcvn(md_file, docx_file)
            converted_count += 1
        except Exception as e:
            print(f"[FAIL] Failed to convert {md_file.name}: {e}")

    print("\n" + "=" * 65)
    print(f"  [DONE] TCVN Conversion Complete: {converted_count}/{len(md_files)} files formatted!")
    print("=" * 65)


if __name__ == "__main__":
    main()
